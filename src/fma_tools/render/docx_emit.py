"""Finished HTML to Word, through a closed vocabulary that refuses what it cannot keep.

The old engine's iron rule was "never build a DOCX by parsing HTML" -- right when every
DOCX was a second, hand-coded authoring of the same numbers. Here the HTML is the
single authored artifact, so a STRICT parser of a closed vocabulary is the anti-drift
mechanism: any element it has not been taught refuses with the line, the element and
the fix, because a lenient parser is how content silently vanishes from the Word copy.

Charts ride as pictures: `svg`, `canvas` and anything marked data-render="image" are
screenshotted off the same browser layout the PDF used, so the Annex bars and KPI
graphics that a hand retype used to lose arrive as the exact pixels the PDF carries.
Tables must declare a `<colgroup>` -- widths are never re-derived, so the two documents
cannot disagree about the same table.
"""

from __future__ import annotations

import base64
import io
import re
from html.parser import HTMLParser
from pathlib import Path

from ..errors import Refusal

_VOID = {"br", "img", "col", "hr", "meta", "link"}
_SKIP_CONTENT = {"style", "script", "title", "head"}
_IMG_LIKE = ("svg", "canvas")

_BLOCK_TAGS = {"html", "body", "h1", "h2", "h3", "h4", "p", "ul", "ol", "li",
               "table", "colgroup", "col", "thead", "tbody", "tr", "th", "td",
               "img", "figure", "figcaption", "hr", "header", "footer", "div",
               "svg", "canvas"}
_INLINE_TAGS = {"strong", "b", "em", "i", "br", "span", "sub", "sup"}
_DIV_CLASSES = {"callout", "kpi-grid", "kpi", "k", "v", "u", "page-break"}

_FONT = "Calibri"


class _Node:
    __slots__ = ("tag", "attrs", "children", "line", "classes")

    def __init__(self, tag, attrs, line):
        self.tag = tag
        self.attrs = dict(attrs)
        self.children = []          # _Node or str
        self.line = line
        self.classes = set((self.attrs.get("class") or "").split())

    def text(self) -> str:
        out = []
        for c in self.children:
            out.append(c if isinstance(c, str) else
                       ("\n" if c.tag == "br" else c.text()))
        return re.sub(r"[ \t\r\f]+", " ", "".join(out)).strip()

    def find_all(self, tag):
        for c in self.children:
            if isinstance(c, _Node):
                if c.tag == tag:
                    yield c
                yield from c.find_all(tag)


def _refuse(line: int, what: str, fix: str) -> Refusal:
    return Refusal("VOCAB_REFUSED",
                   f"REFUSED docx vocabulary at line {line}: {what}. {fix} "
                   "(taught blocks: h1-h4, p, table+colgroup, ul/ol, img[data:], "
                   "figure, div.callout, div.kpi-grid, div.page-break, hr, "
                   "header/footer, svg / data-render=\"image\"; inline: strong, em, "
                   "br, span)")


class _TreeParser(HTMLParser):
    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _Node("root", [], 0)
        self.stack = [self.root]
        self._skip_depth = 0
        self._imglike_depth = 0

    def handle_starttag(self, tag, attrs):
        line = self.getpos()[0]
        attrs_d = dict(attrs)
        if self._imglike_depth:
            if tag not in _VOID:
                self._imglike_depth += 1
            return
        if tag in _SKIP_CONTENT:
            self._skip_depth += 1
            return
        if self._skip_depth:
            return
        if tag in _IMG_LIKE:
            node = _Node(tag, attrs, line)
            self.stack[-1].children.append(node)
            self._imglike_depth = 1
            return
        if tag not in _BLOCK_TAGS and tag not in _INLINE_TAGS:
            raise _refuse(line, f"<{tag}> is not taught",
                          "Rewrite it with taught blocks, or mark the element "
                          "data-render=\"image\" to carry it into Word as a picture.")
        if tag == "div":
            classes = set((attrs_d.get("class") or "").split())
            if "data-render" not in attrs_d and not (classes & _DIV_CLASSES):
                raise _refuse(line, f"<div class=\"{attrs_d.get('class', '')}\"> is "
                              "not a taught block",
                              "Use div.callout, div.kpi-grid or div.page-break, or "
                              "restate the content as headings, paragraphs and "
                              "tables.")
        node = _Node(tag, attrs, line)
        self.stack[-1].children.append(node)
        if "data-render" in attrs_d and attrs_d.get("data-render") == "image":
            self._imglike_depth = 1
            return
        if tag not in _VOID:
            self.stack.append(node)

    def handle_startendtag(self, tag, attrs):
        was_imglike = self._imglike_depth
        self.handle_starttag(tag, attrs)
        if was_imglike:
            # a self-closing tag inside an image node opened nothing to close
            if tag not in _VOID:
                self._imglike_depth -= 1
        elif self._imglike_depth == 1:
            self._imglike_depth = 0          # a self-closing image node has no body
        elif tag not in _VOID and not self._skip_depth \
                and tag not in _SKIP_CONTENT and self.stack[-1].tag == tag:
            self.stack.pop()

    def handle_endtag(self, tag):
        if self._imglike_depth:
            if tag not in _VOID:
                self._imglike_depth -= 1
            return
        if tag in _SKIP_CONTENT:
            self._skip_depth = max(0, self._skip_depth - 1)
            return
        if self._skip_depth or tag in _VOID:
            return
        while len(self.stack) > 1:
            top = self.stack.pop()
            if top.tag == tag:
                return

    def handle_data(self, data):
        if self._skip_depth or self._imglike_depth:
            return
        if data.strip():
            self.stack[-1].children.append(data)


def _is_imagey(node: _Node) -> bool:
    return node.tag in _IMG_LIKE or node.attrs.get("data-render") == "image"


class DocxEmitter:
    """Walk the vocabulary tree into python-docx. Returns the parity ledger."""

    def __init__(self, geometry, shots: list[str]):
        from docx import Document
        from docx.shared import Mm, Pt
        self.doc = Document()
        sec = self.doc.sections[0]
        sec.page_width = Mm(geometry.width_pt / 72.0 * 25.4)
        sec.page_height = Mm(geometry.height_pt / 72.0 * 25.4)
        sec.left_margin = sec.right_margin = Mm(20)
        sec.top_margin = sec.bottom_margin = Mm(18)
        st = self.doc.styles["Normal"]
        st.font.name = _FONT
        st.font.size = Pt(10)
        self.shots = shots
        self.shot_i = 0
        self.expected_texts: list[str] = []
        self.expected_pictures = 0

    # -- inline ---------------------------------------------------------------
    def _runs_into(self, para, node: _Node, bold=False, italic=False, colour=None):
        from docx.shared import RGBColor
        for c in node.children:
            if isinstance(c, str):
                text = re.sub(r"[ \t\r\f]+", " ", c)
                if not text:
                    continue
                run = para.add_run(text)
                run.font.name = _FONT
                run.bold = bold or None
                run.italic = italic or None
                if colour:
                    run.font.color.rgb = RGBColor(*colour)
            elif c.tag == "br":
                para.add_run().add_break()
            elif c.tag in ("strong", "b"):
                self._runs_into(para, c, True, italic, colour)
            elif c.tag in ("em", "i"):
                self._runs_into(para, c, bold, True, colour)
            elif c.tag in ("span", "sub", "sup"):
                col = (0x6B, 0x72, 0x80) if "muted" in c.classes else colour
                self._runs_into(para, c, bold, italic, col)
            elif c.tag == "img":
                self._picture_run(para, c)
            elif _is_imagey(c):
                self._shot_run(para, c)
            else:
                raise _refuse(c.line, f"<{c.tag}> inside a paragraph is not taught",
                              "Only strong, em, br, span and img are inline.")

    def _record(self, para):
        text = para.text.strip()
        if text:
            self.expected_texts.append(re.sub(r"\s+", " ", text))

    def _para(self, node: _Node, size, bold=False, colour=None, space_before=6,
              space_after=4):
        from docx.shared import Pt, RGBColor
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        self._runs_into(p, node, bold=bold,
                        colour=colour)
        for run in p.runs:
            if run.font.size is None:
                run.font.size = Pt(size)
        if colour:
            for run in p.runs:
                if run.font.color and run.font.color.rgb is None:
                    run.font.color.rgb = RGBColor(*colour)
        self._record(p)
        return p

    def _picture_run(self, para, node: _Node):
        from docx.shared import Inches
        src = node.attrs.get("src", "")
        m = re.match(r"data:image/(png|jpeg|gif);base64,(.+)", src, re.S)
        if not m:
            raise _refuse(node.line, "<img> src is not a data: URI",
                          "Embed the image as data:image/png;base64,... -- render is "
                          "network-sealed and a file path breaks on another Mac.")
        blob = base64.b64decode(m.group(2))
        width = node.attrs.get("width")
        run = para.add_run()
        run.add_picture(io.BytesIO(blob),
                        width=Inches(int(width) / 96.0) if width else None)
        self.expected_pictures += 1

    def _shot_run(self, para, node: _Node):
        if self.shot_i >= len(self.shots):
            raise Refusal("GATE_FAILED",
                          f"line {node.line}: a chart node has no matching browser "
                          "screenshot; the HTML and the rendered page diverged")
        from docx.shared import Mm
        run = para.add_run()
        run.add_picture(self.shots[self.shot_i], width=Mm(160))
        self.shot_i += 1
        self.expected_pictures += 1

    # -- blocks ---------------------------------------------------------------
    def walk(self, node: _Node):
        for c in node.children:
            if isinstance(c, str):
                if c.strip():
                    raise _refuse(node.line, f"bare text {c.strip()[:40]!r} outside "
                                  "any block", "Wrap it in <p>.")
                continue
            self.block(c)

    def block(self, node: _Node):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        t = node.tag
        if t in ("html", "body"):
            self.walk(node)
        elif t in ("header", "footer"):
            self._header_footer(node)
        elif t == "h1":
            self._para(node, 20, bold=True, colour=(0x0D, 0x2B, 0x4E),
                       space_before=0, space_after=10)
        elif t == "h2":
            if "annex" in node.classes:
                self.doc.add_page_break()
            badge = node.attrs.get("data-badge")
            p = self._para(node, 14, bold=True, colour=(0x0D, 0x2B, 0x4E),
                           space_before=14, space_after=6)
            if badge and p.runs:
                p.runs[0].text = f"{badge}  {p.runs[0].text}"
                self.expected_texts[-1] = re.sub(r"\s+", " ", p.text.strip())
        elif t == "h3":
            self._para(node, 11.5, bold=True, space_before=10)
        elif t == "h4":
            self._para(node, 10.5, bold=True, space_before=8)
        elif t == "p":
            self._para(node, 10)
        elif t in ("ul", "ol"):
            style = "List Bullet" if t == "ul" else "List Number"
            for li in node.children:
                if isinstance(li, str):
                    continue
                if li.tag != "li":
                    raise _refuse(li.line, f"<{li.tag}> inside a list", "Only <li>.")
                p = self.doc.add_paragraph(style=style)
                self._runs_into(p, li)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)
                self._record(p)
        elif t == "table":
            self._table(node)
        elif t == "figure":
            for c in node.children:
                if isinstance(c, str):
                    continue
                if c.tag == "img" or _is_imagey(c):
                    p = self.doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if c.tag == "img":
                        self._picture_run(p, c)
                    else:
                        self._shot_run(p, c)
                elif c.tag == "figcaption":
                    self._para(c, 8.5, colour=(0x6B, 0x72, 0x80), space_before=2)
                else:
                    raise _refuse(c.line, f"<{c.tag}> inside a figure",
                                  "A figure holds an image and a figcaption.")
        elif t == "img":
            p = self.doc.add_paragraph()
            self._picture_run(p, node)
        elif _is_imagey(node):
            p = self.doc.add_paragraph()
            self._shot_run(p, node)
        elif t == "hr":
            self.doc.add_paragraph()
        elif t == "div" and "page-break" in node.classes:
            self.doc.add_page_break()
        elif t == "div" and "callout" in node.classes:
            self._callout(node)
        elif t == "div" and "kpi-grid" in node.classes:
            self._kpi_grid(node)
        else:
            raise _refuse(node.line, f"<{t} class=\"{' '.join(node.classes)}\"> as a "
                          "block", "Not a taught block here.")

    def _header_footer(self, node: _Node):
        from docx.shared import Pt
        sec = self.doc.sections[0]
        target = sec.header if node.tag == "header" else sec.footer
        p = target.paragraphs[0]
        for c in node.children:
            if isinstance(c, str):
                if c.strip():
                    p.add_run(re.sub(r"\s+", " ", c)).font.size = Pt(8)
            elif c.tag == "span" and "page-number" in c.classes:
                self._page_field(p)
            elif c.tag in _INLINE_TAGS:
                before = len(p.runs)
                self._runs_into(p, c)
                for run in p.runs[before:]:
                    run.font.size = Pt(8)
            else:
                raise _refuse(c.line, f"<{c.tag}> in a header/footer",
                              "Headers hold inline text and span.page-number only.")

    @staticmethod
    def _page_field(p):
        from docx.oxml.ns import qn
        run = p.add_run()
        for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
            if kind:
                el = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): kind})
            else:
                el = run._element.makeelement(qn("w:instrText"), {})
                el.text = text
            run._element.append(el)

    def _table(self, node: _Node):
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm, Pt, RGBColor
        cols = [c for cg in node.find_all("colgroup") for c in cg.children
                if isinstance(c, _Node) and c.tag == "col"]
        widths = []
        for c in cols:
            m = re.search(r"width\s*:\s*([\d.]+)%", c.attrs.get("style", ""))
            if m:
                widths.append(float(m.group(1)))
        rows = list(node.find_all("tr"))
        if not rows:
            raise _refuse(node.line, "a table with no rows", "Add thead/tbody rows.")
        ncols = max(sum(1 for c in r.children
                        if isinstance(c, _Node) and c.tag in ("th", "td"))
                    for r in rows)
        if len(widths) != ncols:
            raise _refuse(node.line,
                          f"table declares {len(widths)} colgroup widths for "
                          f"{ncols} columns",
                          "Every table needs <colgroup><col style=\"width:NN%\">"
                          " per column -- widths are never re-derived, so the PDF "
                          "and the Word file cannot disagree about the same table.")
        sec = self.doc.sections[0]
        usable = sec.page_width - sec.left_margin - sec.right_margin
        tbl = self.doc.add_table(rows=0, cols=ncols)
        tbl.style = "Table Grid"
        tbl.autofit = False
        for r in rows:
            cells = [c for c in r.children
                     if isinstance(c, _Node) and c.tag in ("th", "td")]
            for c in cells:
                if "colspan" in c.attrs or "rowspan" in c.attrs:
                    raise _refuse(c.line, "colspan/rowspan is not taught",
                                  "A spanned cell shifts every later figure under "
                                  "the wrong column heading in Word; restructure "
                                  "the table into plain rows.")
            row = tbl.add_row()
            is_total = "total" in r.classes
            for i in range(ncols):
                cell = row.cells[i]
                cell.width = int(usable * widths[i] / 100)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if i >= len(cells):
                    continue
                spec = cells[i]
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                self._runs_into(p, spec)
                head = spec.tag == "th"
                for run in p.runs:
                    run.font.name = _FONT
                    if run.font.size is None:
                        run.font.size = Pt(8.5 if head else 9.5)
                    if head:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
                    elif is_total:
                        run.bold = True
                if "num" in spec.classes:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                self._record(p)

    def _callout(self, node: _Node):
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.shared import Pt
        tbl = self.doc.add_table(rows=1, cols=1)
        cell = tbl.rows[0].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shd = cell._tc.get_or_add_tcPr().makeelement(
            qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "F3F4F6"})
        cell._tc.get_or_add_tcPr().append(shd)
        label = node.attrs.get("data-label")
        first = cell.paragraphs[0]
        if label:
            run = first.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            self._record(first)
        for c in node.children:
            if isinstance(c, str):
                if c.strip():
                    p = cell.add_paragraph()
                    p.add_run(re.sub(r"\s+", " ", c).strip()).font.size = Pt(10)
                    self._record(p)
            elif c.tag == "p":
                p = cell.add_paragraph()
                self._runs_into(p, c)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(10)
                self._record(p)
            elif c.tag in ("ul", "ol"):
                for li in c.children:
                    if isinstance(li, _Node) and li.tag == "li":
                        p = cell.add_paragraph(style="List Bullet")
                        self._runs_into(p, li)
                        self._record(p)
            else:
                raise _refuse(getattr(c, "line", node.line),
                              "a callout holds paragraphs and lists only",
                              "Move other blocks outside the callout.")

    def _kpi_grid(self, node: _Node):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
        kpis = [c for c in node.children
                if isinstance(c, _Node) and "kpi" in c.classes]
        if not kpis:
            raise _refuse(node.line, "a kpi-grid with no div.kpi tiles",
                          "Each tile is div.kpi holding .k, .v and .u.")
        per_row = 3
        tbl = self.doc.add_table(rows=0, cols=per_row)
        tbl.style = "Table Grid"
        for i in range(0, len(kpis), per_row):
            row = tbl.add_row()
            for j, kpi in enumerate(kpis[i:i + per_row]):
                cell = row.cells[j]
                parts = {}
                for n in kpi.children:
                    if isinstance(n, str):
                        if n.strip():
                            raise _refuse(kpi.line, "bare text inside a kpi tile",
                                          "A tile holds .k, .v and .u elements only.")
                        continue
                    tile_cls = n.classes & {"k", "v", "u"}
                    if len(tile_cls) != 1:
                        raise _refuse(n.line, f"<{n.tag} class=\"{' '.join(n.classes)}\">"
                                      " inside a kpi tile",
                                      "Each tile child carries exactly one of "
                                      ".k, .v, .u -- anything else would be "
                                      "silently dropped.")
                    cls = tile_cls.pop()
                    if cls in parts:
                        raise _refuse(n.line, f"duplicate .{cls} in one kpi tile",
                                      "One .k, one .v, one .u per tile.")
                    parts[cls] = n
                for cls, size, bold, colour in (("k", 8, True, (0x6B, 0x72, 0x80)),
                                                ("v", 14, True, (0x0D, 0x2B, 0x4E)),
                                                ("u", 8, False, (0x6B, 0x72, 0x80))):
                    n = parts.get(cls)
                    if n is None:
                        continue
                    p = cell.paragraphs[0] if cls == "k" else cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(n.text())
                    run.font.size = Pt(size)
                    run.bold = bold
                    run.font.color.rgb = RGBColor(*colour)
                    self._record(p)


def count_chart_nodes(html: str) -> int:
    """svg / canvas / data-render="image" nodes needing a browser screenshot.
    <img data:> does not count -- it decodes without a browser."""
    n = len(re.findall(r"<svg\b", html)) + len(re.findall(r"<canvas\b", html))
    n += len(re.findall(r"data-render=\"image\"", html))
    return n


def emit(html: str, geometry, shots: list[str], out_path: Path) -> dict:
    parser = _TreeParser()
    parser.feed(html)
    em = DocxEmitter(geometry, shots)
    em.walk(parser.root)
    em.doc.save(str(out_path))
    return {"expected_texts": em.expected_texts,
            "expected_pictures": em.expected_pictures}


def saved_texts_and_pictures(path: Path) -> tuple[list[str], int]:
    """Body paragraph and table-cell texts plus picture count, off the SAVED file."""
    from docx import Document
    doc = Document(str(path))
    texts = [re.sub(r"\s+", " ", p.text.strip())
             for p in doc.paragraphs if p.text.strip()]

    def _tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            texts.append(re.sub(r"\s+", " ", p.text.strip()))
                    _tables(cell.tables)

    _tables(doc.tables)
    return texts, len(doc.inline_shapes)
